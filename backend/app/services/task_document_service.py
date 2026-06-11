import os
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.config import settings


def _safe_filename(text: str):
    keep = "".join(c for c in text if c.isalnum() or c in (" ", "-", "_")).strip()
    return keep[:80].replace(" ", "_") or "task"


def create_task_docx(task, update_link: str, sender_name: str, sender_designation: str):
    os.makedirs(settings.TASKS_DIR, exist_ok=True)

    filename = f"task_{task.id}_{_safe_filename(task.title)}.docx"
    path = os.path.join(settings.TASKS_DIR, filename)

    doc = Document()
    doc.add_heading("Task Assignment Document", level=1)

    doc.add_paragraph(f"Task Title: {task.title}")
    doc.add_paragraph(f"Priority: {task.priority.title()}")
    doc.add_paragraph(f"Deadline: {task.due_date or 'Not specified'}")
    doc.add_paragraph(f"Assigned To: {task.assignees or task.assigned_to or 'Not specified'}")
    doc.add_paragraph(f"Generated On: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")

    doc.add_heading("Task Details", level=2)
    doc.add_paragraph(task.professional_description or task.description or "")

    doc.add_heading("Task Update Link", level=2)
    doc.add_paragraph(update_link)

    doc.add_heading("Instructions", level=2)
    doc.add_paragraph(
        "Please open the task update link, complete checklist items, update progress, "
        "add comments if required, and upload proof of work when applicable."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph(sender_name)
    doc.add_paragraph(sender_designation)

    doc.save(path)
    return path


def create_task_pdf(task, update_link: str, sender_name: str, sender_designation: str):
    os.makedirs(settings.TASKS_DIR, exist_ok=True)

    filename = f"task_{task.id}_{_safe_filename(task.title)}.pdf"
    path = os.path.join(settings.TASKS_DIR, filename)

    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4

    y = height - 50

    def write_line(text, size=10, gap=16):
        nonlocal y
        if y < 60:
            c.showPage()
            y = height - 50

        c.setFont("Helvetica", size)
        c.drawString(50, y, str(text)[:110])
        y -= gap

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Task Assignment Document")
    y -= 35

    write_line(f"Task Title: {task.title}", 11)
    write_line(f"Priority: {task.priority.title()}", 11)
    write_line(f"Deadline: {task.due_date or 'Not specified'}", 11)
    write_line(f"Assigned To: {task.assignees or task.assigned_to or 'Not specified'}", 11)
    write_line(f"Generated On: {datetime.now().strftime('%d %b %Y, %I:%M %p')}", 11)

    y -= 10
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "Task Details")
    y -= 22

    description = task.professional_description or task.description or ""
    for paragraph in description.splitlines():
        if not paragraph.strip():
            y -= 8
            continue
        write_line(paragraph.strip(), 10, 15)

    y -= 10
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "Task Update Link")
    y -= 22

    write_line(update_link, 10, 15)

    y -= 10
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "Instructions")
    y -= 22

    instruction = (
        "Please open the task update link, complete checklist items, update progress, "
        "add comments if required, and upload proof of work when applicable."
    )

    for chunk in [instruction[i:i + 100] for i in range(0, len(instruction), 100)]:
        write_line(chunk, 10, 15)

    y -= 20
    write_line("Best regards,", 10)
    write_line(sender_name, 10)
    write_line(sender_designation, 10)

    c.save()
    return path