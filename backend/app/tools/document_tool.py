"""
DOCX export tool using python-docx.
"""
import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings


def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def export_to_docx(content: str, doc_type: str, title: str = "") -> str:
    """
    Create a DOCX file from content string.
    Returns the absolute file path.
    """
    _ensure_dir(settings.OUTPUT_DIR)

    # Pick output folder
    folder_map = {
        "email": settings.EMAILS_DIR,
        "report": settings.REPORTS_DIR,
        "meeting": settings.MOMS_DIR,
    }
    folder = folder_map.get(doc_type, settings.OUTPUT_DIR)
    _ensure_dir(folder)

    filename = f"{doc_type}_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(folder, filename)

    doc = Document()

    # Title
    heading = doc.add_heading(title or doc_type.upper(), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x53, 0x8A)

    # Timestamp
    ts = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ts.runs[0].font.size = Pt(9)
    ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Content – split by newlines
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(filepath)
    return filepath
